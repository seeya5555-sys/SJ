"""
Boarding Report — Word(.docx) 생성

방선보고서 + Defect List 통합 양식:
  · 표지: "Vessel boarding report" 헤더 + 결재란 + 정보 표 + Master/CE
  · 본문: 섹션별 블록 (paragraph / bullet / table / image / info_table / defect_table)
  · Sinokor 푸터: "CODE<107-301>/2015.04.17 ... Sinokor Ship Management Co., Ltd"

Dry Dock 모듈(dock_report_docx.py)의 헬퍼 일부를 임포트해 재사용.
"""
import io
import json
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Dry Dock 모듈에서 검증된 헬퍼 재사용
from dock_report_docx import (
    _set_cell_shading, _set_cell_borders, _set_font, _add_paragraph,
    _set_row_height, _set_table_fixed_layout, _add_horizontal_line,
    _crop_to_aspect, _GLOBAL_TEMP_FILES,
)


# ─────────────────────────────────────────────────────────────
#  방선보고서 표지 — "Vessel boarding report" 헤더 + 결재란 + 정보 표
# ─────────────────────────────────────────────────────────────
def _build_brep_cover(doc, report):
    """
    회사 양식 (260331_Maritime_Glory호_방선보고서.docx) 기준 - 사용자 보정:
      · 결재자 4칸 균등 너비
      · Inspector 행 중복 제거 (1줄만)
      · 9컬럼 × 6행 통합 표
    """
    n_cols = 9
    n_rows = 6   # 헤더 + 좌측 정보 3행(Vessel/Port/Inspector) + Date행 + Master/CE행
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 컬럼 너비 (총 17.68cm)
    #   col 0: 정보 라벨   2.42
    #   col 1: 정보 값     4.46
    #   col 2: "결재" 라벨  0.71
    #   col 3~8: 결재자 6칸 → 4명 균등 분배
    #     기안자=2칸, 팀장=2칸, 중역=1칸, 대표이사=1칸 (양식과 동일)
    #     단, 각 결재자 총 너비를 균등하게: 17.68 - 2.42 - 4.46 - 0.71 = 10.09cm
    #     4명으로 나누면 한 명당 ≈ 2.52cm
    #     기안자(2칸) = 1.26+1.26 = 2.52
    #     팀장(2칸)   = 1.26+1.26 = 2.52
    #     중역(1칸)   = 2.52
    #     대표이사(1칸) = 2.53 (소수점 보정)
    col_cm = [2.42, 4.46, 0.71,
              1.26, 1.26,    # 기안자 (2칸)
              1.26, 1.26,    # 팀장 (2칸)
              2.52,          # 중역
              2.53]          # 대표이사
    total_cm = sum(col_cm)

    # ─── R0: "Vessel Boarding Report" 전체 병합 헤더 ───
    hcell = tbl.rows[0].cells[0]
    for ci in range(1, n_cols):
        hcell = hcell.merge(tbl.rows[0].cells[ci])
    _set_cell_borders(hcell)
    hcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    hp = hcell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run('Vessel Boarding Report')
    _set_font(hr, size=14, bold=True)
    _set_row_height(tbl.rows[0], 0.8)

    # ─── 좌측 정보 영역 (R1~R3: col 0, 1) — Inspector 1행만 ───
    left_info = [
        ('Vessel',    report.get('vessel_name') or ''),
        ('Port',      report.get('port') or ''),
        ('Inspector', report.get('supervisor_name') or ''),
    ]
    for ri_offset, (label, value) in enumerate(left_info):
        ri = 1 + ri_offset
        lc = tbl.rows[ri].cells[0]
        _set_cell_shading(lc, 'F2F2F2')
        _set_cell_borders(lc)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lr = lp.add_run(label)
        _set_font(lr, size=10, bold=True)

        vc = tbl.rows[ri].cells[1]
        _set_cell_borders(vc)
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vr = vp.add_run(str(value))
        _set_font(vr, size=10)
        _set_row_height(tbl.rows[ri], 0.65)

    # ─── 결재 영역 (R1~R3: col 2~8) ───
    # col 2: "결재" 세로 라벨 (R1~R3 병합)
    app_label = tbl.rows[1].cells[2]
    for ri in range(2, 4):
        app_label = app_label.merge(tbl.rows[ri].cells[2])
    _set_cell_shading(app_label, 'F2F2F2')
    _set_cell_borders(app_label)
    app_label.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    alp = app_label.paragraphs[0]
    alp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    alr = alp.add_run('결\n재')
    _set_font(alr, size=10, bold=True)

    # 결재자 헤더 (R1) + 서명란 (R2) + "/" 행 (R3)
    approver_spans = [
        ('기 안 자', 3, 4),
        ('팀 장',    5, 6),
        ('중 역',    7, 7),
        ('대표이사', 8, 8),
    ]
    approver_values = [
        report.get('approval_drafter')   or '',
        report.get('approval_team_lead') or '',
        report.get('approval_director')  or '',
        report.get('approval_ceo')       or '',
    ]

    # R1: 헤더
    for label, c_start, c_end in approver_spans:
        cell = tbl.rows[1].cells[c_start]
        for ci in range(c_start + 1, c_end + 1):
            cell = cell.merge(tbl.rows[1].cells[ci])
        _set_cell_shading(cell, 'F2F2F2')
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        _set_font(r, size=9.5, bold=True)

    # R2: 서명 (이름)
    for k, (_, c_start, c_end) in enumerate(approver_spans):
        cell = tbl.rows[2].cells[c_start]
        for ci in range(c_start + 1, c_end + 1):
            cell = cell.merge(tbl.rows[2].cells[ci])
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if approver_values[k]:
            r = p.add_run(approver_values[k])
            _set_font(r, size=10)
    _set_row_height(tbl.rows[2], 1.1)

    # R3: "/" 행 (서명 날짜용)
    for _, c_start, c_end in approver_spans:
        cell = tbl.rows[3].cells[c_start]
        for ci in range(c_start + 1, c_end + 1):
            cell = cell.merge(tbl.rows[3].cells[ci])
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('/')
        _set_font(r, size=10)

    # ─── R4: Date/time | 값(5병합) | Ship-Visit Score(2병합) | 값 ───
    c = tbl.rows[4].cells[0]
    _set_cell_shading(c, 'F2F2F2'); _set_cell_borders(c)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('Date/time'); _set_font(r, size=10, bold=True)

    vc = tbl.rows[4].cells[1]
    for ci in range(2, 6):
        vc = vc.merge(tbl.rows[4].cells[ci])
    _set_cell_borders(vc)
    vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    vp = vc.paragraphs[0]; vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run(_fmt_period(report.get('boarding_start'),
                                report.get('boarding_end')))
    _set_font(vr, size=10)

    sc = tbl.rows[4].cells[6]
    sc = sc.merge(tbl.rows[4].cells[7])
    _set_cell_shading(sc, 'F2F2F2'); _set_cell_borders(sc)
    sc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    sp = sc.paragraphs[0]; sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run('Ship-Visit\nChecklist Score'); _set_font(sr, size=9.5, bold=True)

    sv = tbl.rows[4].cells[8]
    _set_cell_borders(sv)
    sv.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    svp = sv.paragraphs[0]; svp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    svr = svp.add_run(report.get('sv_checklist_score') or '-')
    _set_font(svr, size=10)
    _set_row_height(tbl.rows[4], 0.85)

    # ─── R5: Master(boarding date) | 값(3병합) | C/E(boarding date)(2병합) | 값(3병합) ───
    mc = tbl.rows[5].cells[0]
    _set_cell_shading(mc, 'F2F2F2'); _set_cell_borders(mc)
    mc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    mp = mc.paragraphs[0]; mp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    mr = mp.add_run('Master\n(boarding date)'); _set_font(mr, size=9.5, bold=True)

    mv = tbl.rows[5].cells[1]
    for ci in range(2, 4):
        mv = mv.merge(tbl.rows[5].cells[ci])
    _set_cell_borders(mv)
    mv.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    mvp = mv.paragraphs[0]; mvp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mvr = mvp.add_run(_fmt_person_date(report.get('master_name'),
                                       report.get('master_board_date')))
    _set_font(mvr, size=10)

    cc = tbl.rows[5].cells[4]
    cc = cc.merge(tbl.rows[5].cells[5])
    _set_cell_shading(cc, 'F2F2F2'); _set_cell_borders(cc)
    cc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cp = cc.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run('C/E\n(boarding date)'); _set_font(cr, size=9.5, bold=True)

    cv = tbl.rows[5].cells[6]
    for ci in range(7, 9):
        cv = cv.merge(tbl.rows[5].cells[ci])
    _set_cell_borders(cv)
    cv.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cvp = cv.paragraphs[0]; cvp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cvr = cvp.add_run(_fmt_person_date(report.get('chief_eng_name'),
                                       report.get('chief_eng_board_date')))
    _set_font(cvr, size=10)
    _set_row_height(tbl.rows[5], 0.85)

    # 표 전체 너비 고정
    _set_table_fixed_layout(tbl, total_cm, col_cm)
    _add_paragraph(doc, '', before=0, after=2)


def _fmt_period(start, end):
    if not start and not end:
        return ''
    s = (start or '').replace('-', '.')
    e = (end or '').replace('-', '.')
    if s and e:
        return f'{s} ~ {e}'
    return s or e


def _fmt_person_date(name, date):
    name = (name or '').strip()
    date = (date or '').replace('-', '.')
    if name and date:
        return f'{name}  ({date})'
    return name or date or ''


# ─────────────────────────────────────────────────────────────
#  본문 — 섹션 + 블록
# ─────────────────────────────────────────────────────────────
def _render_brep_sections(doc, sections_tree, depth=0, prefix=''):
    """
    1단계 섹션은 회사 양식 그대로:
      [ 섹션 제목 (좌측 라벨) | 본문 블록들 (우측 큰 셀) ]
    2단계 이하는 일반 제목 + 블록 형태 (우측 셀 내부에 들여쓰기)
    """
    for i, sec in enumerate(sections_tree):
        num = f'{prefix}-{i + 1}' if prefix else f'{i + 1}'

        if depth == 0:
            # 회사 양식: 큰 통합 표 (좌측 라벨 | 우측 본문)
            _render_brep_section_as_box(doc, sec, num)
        else:
            # 2단계 이하: 양식엔 없지만 사용자가 추가한 하위 섹션
            # 우측 셀 안에서 동작하므로 단순 제목으로 표시
            if depth == 1:
                _add_paragraph(doc, f'■ {sec["title"]}',
                               size=11.5, bold=True, color='1F4E79',
                               before=8, after=4)
            else:
                _add_paragraph(doc, f'· {sec["title"]}',
                               size=11, bold=True,
                               before=6, after=3)
            blocks = sorted(sec.get('blocks', []),
                            key=lambda b: (b.get('display_order', 0), b.get('id', 0)))
            for b in blocks:
                _render_brep_block(doc, b, depth)
            if sec.get('children'):
                _render_brep_sections(doc, sec['children'], depth + 1, num)


def _render_brep_section_as_box(doc, sec, num):
    """
    1단계 섹션을 회사 양식의 큰 통합 표로 렌더링.
      · 행 1개, 열 2개
      · 좌측: 섹션 제목 (라벨, 세로 가운데)
      · 우측: 모든 블록 콘텐츠
    """
    # 섹션 제목에서 자동 번호 prefix("1. ") 제거 — 회사 양식엔 번호가 없음
    title = (sec.get('title') or '').strip()
    # "1. " "1) " 같은 앞부분 자동 제거
    import re
    clean_title = re.sub(r'^\s*\d+[.)]\s*', '', title)

    # 표 1행 × 2열
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_cm = [2.6, 14.4]   # 합계 17.0cm (표지와 동일)
    total_cm = sum(col_cm)

    # 좌측 라벨 셀
    lcell = tbl.rows[0].cells[0]
    _set_cell_shading(lcell, 'F2F2F2')
    _set_cell_borders(lcell)
    lcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = lcell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run(clean_title)
    _set_font(lr, size=11, bold=True)

    # 우측 본문 셀
    rcell = tbl.rows[0].cells[1]
    _set_cell_borders(rcell)
    rcell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    # 기본 paragraph 제거 (빈 줄 방지)
    p0 = rcell.paragraphs[0]
    p0._element.getparent().remove(p0._element)

    # 블록들을 우측 셀에 렌더링
    blocks = sorted(sec.get('blocks', []),
                    key=lambda b: (b.get('display_order', 0), b.get('id', 0)))
    for b in blocks:
        _render_brep_block_in_cell(rcell, b)

    # 자식 섹션도 우측 셀 안에
    if sec.get('children'):
        for ci, child in enumerate(sec['children']):
            child_title = re.sub(r'^\s*\d+[.)]\s*', '',
                                 (child.get('title') or '').strip())
            cp = rcell.add_paragraph()
            cp.paragraph_format.space_before = Pt(8)
            cp.paragraph_format.space_after = Pt(4)
            cr = cp.add_run(f'■ {child_title}')
            _set_font(cr, size=11.5, bold=True, color='1F4E79')
            for b in sorted(child.get('blocks', []),
                            key=lambda b: (b.get('display_order', 0), b.get('id', 0))):
                _render_brep_block_in_cell(rcell, b)

    _set_table_fixed_layout(tbl, total_cm, col_cm)

    # 표와 표 사이 약간의 공백
    _add_paragraph(doc, '', before=0, after=4)


def _render_brep_block_in_cell(cell, block):
    """우측 본문 셀 안에 블록 하나 렌더링.
    info_table/defect_table은 통합 표 바깥으로 빼고, 일반 블록만 안에 넣음.
    """
    bt = block.get('block_type')
    content = block.get('content') or {}
    if isinstance(content, str):
        try: content = json.loads(content)
        except Exception: content = {}

    if bt == 'paragraph':
        text = (content.get('text') or '').strip()
        if not text:
            return
        for line in text.split('\n'):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(line)
            _set_font(r, size=10.5)

    elif bt == 'bullet_list':
        items = content.get('items') or []
        marker = content.get('marker') or 'bullet'
        counters = [0, 0, 0, 0]
        for it in items:
            if isinstance(it, str):
                text, indent = it, 0
            else:
                text = it.get('text', '')
                indent = max(0, min(3, it.get('indent', 0)))
            if not text.strip():
                continue
            counters[indent] += 1
            for k in range(indent + 1, 4):
                counters[k] = 0
            n = counters[indent]
            if marker == 'dash':      mk = '–'
            elif marker == 'number':  mk = _number_by_depth(indent, n)
            elif marker == 'alpha':   mk = _alpha_by_depth(indent, n)
            else:                      mk = '•'

            # 항목 내 줄바꿈(\n) — 둘째 줄부터 마커 폭만큼 추가 들여쓰기
            base_left = 0.5 * indent
            cont_left = base_left + 0.5
            lines = text.split('\n')
            for li, ln in enumerate(lines):
                p = cell.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                if li == 0:
                    p.paragraph_format.left_indent = Cm(base_left)
                    r = p.add_run(f'{mk}  {ln}')
                else:
                    p.paragraph_format.left_indent = Cm(cont_left)
                    r = p.add_run(ln)
                _set_font(r, size=10.5)

    elif bt == 'table':
        # 일반 표 — 셀 내부에 표 중첩 (LibreOffice 호환 제한 있어서 단순 형태로)
        _render_table_in_cell(cell, content)

    elif bt == 'image':
        _render_image_in_cell(cell, content)

    elif bt == 'info_table':
        # 셀 내부에 info_table을 단순 라벨/값 paragraph로 표현
        rows = content.get('rows') or []
        for r in rows:
            if not (r.get('label') or '').strip() and not (r.get('value') or '').strip():
                continue
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            lr = p.add_run((r.get('label') or '') + ': ')
            _set_font(lr, size=10.5, bold=True)
            vr = p.add_run(r.get('value') or '')
            _set_font(vr, size=10.5)

    elif bt == 'defect_table':
        # defect_table은 통합 표 안에 넣기엔 너무 복잡 → 안내 문구만
        p = cell.add_paragraph()
        r = p.add_run('※ Defect List는 다음 표를 참조하세요.')
        _set_font(r, size=9.5, color='6B7280')
        r.italic = True


def _render_table_in_cell(parent_cell, content):
    """우측 본문 셀 안에 일반 표 중첩 — 셀 너비 좁으니 단순화"""
    headers = content.get('headers') or []
    rows = content.get('rows') or []
    if not headers and not rows:
        return
    n_cols = max(len(headers), max((len(r) for r in rows), default=0))
    if n_cols == 0:
        return
    headers = (list(headers) + [''] * n_cols)[:n_cols]
    rows = [(list(r) + [''] * n_cols)[:n_cols] for r in rows]

    tbl = parent_cell.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.autofit = False

    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        _set_cell_shading(c, 'D9E2EC')
        _set_cell_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        _set_font(r, size=9.5, bold=True)

    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = tbl.rows[ri].cells[ci]
            _set_cell_borders(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lines = str(val or '').split('\n')
            for li, ln in enumerate(lines):
                if li > 0:
                    p = c.add_paragraph()
                r = p.add_run(ln)
                _set_font(r, size=9.5)

    # 셀 내부 표는 너비 제한 (우측 셀이 약 14cm)
    col_cm_in = [14.0 / n_cols] * n_cols
    _set_table_fixed_layout(tbl, 14.0, col_cm_in)


def _render_image_in_cell(parent_cell, content):
    """우측 본문 셀 안에 이미지 갤러리 (셀 내부 표 중첩)"""
    images = content.get('images') or []
    columns = max(1, min(4, int(content.get('columns', 2) or 2)))
    if not images:
        return

    n_rows = (len(images) + columns - 1) // columns
    cell_total = 14.0   # 우측 본문 셀 가용 폭
    cell_cm = cell_total / columns
    img_w = cell_cm - 0.4
    img_h = img_w * 3 / 4

    tbl = parent_cell.add_table(rows=n_rows * 2, cols=columns)
    tbl.autofit = False

    for idx, img in enumerate(images):
        ri = (idx // columns) * 2
        ci = idx % columns
        img_cell = tbl.rows[ri].cells[ci]
        cap_cell = tbl.rows[ri + 1].cells[ci]
        # 회사 양식은 사진 표에 굵은 테두리 없음 → 옅게
        _set_cell_borders(img_cell, color='D1D5DB')
        _set_cell_borders(cap_cell, color='D1D5DB')

        img_path = _resolve_brep_image_path(img.get('url') or '',
                                            img.get('filename') or '')
        p = img_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if img_path and os.path.exists(img_path):
            try:
                processed = _crop_to_aspect(img_path, 4/3)
                if processed != img_path:
                    _GLOBAL_TEMP_FILES.append(processed)
                run = p.add_run()
                run.add_picture(processed, width=Cm(img_w), height=Cm(img_h))
            except Exception:
                r = p.add_run('[이미지 오류]')
                _set_font(r, size=9, color='B91C1C')
        else:
            r = p.add_run('[이미지 없음]')
            _set_font(r, size=9, color='9CA3AF')

        cap_p = cap_cell.paragraphs[0]
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = img.get('caption') or ''
        if caption:
            cr = cap_p.add_run(caption)
            _set_font(cr, size=9, color='4B5563')
            cr.italic = True

    col_cm_list = [cell_cm] * columns
    _set_table_fixed_layout(tbl, cell_total, col_cm_list)


def _render_brep_block(doc, block, depth):
    bt = block.get('block_type')
    content = block.get('content') or {}
    if isinstance(content, str):
        try: content = json.loads(content)
        except Exception: content = {}

    base_indent = 0.3 + 0.3 * depth

    if bt == 'paragraph':
        _render_paragraph(doc, content, base_indent)
    elif bt == 'bullet_list':
        _render_bullet(doc, content, base_indent)
    elif bt == 'table':
        _render_table(doc, content, base_indent)
    elif bt == 'image':
        _render_image(doc, content, base_indent)
    elif bt == 'info_table':
        _render_info_table(doc, content, base_indent)
    elif bt == 'defect_table':
        _render_defect_table(doc, content, base_indent)


def _render_paragraph(doc, content, base_indent):
    text = (content.get('text') or '').strip()
    if not text:
        return
    for line in text.split('\n'):
        _add_paragraph(doc, line, size=10.5,
                       indent_left=base_indent, after=4)


# 마커 형식: 깊이별
CIRCLED = ['①','②','③','④','⑤','⑥','⑦','⑧','⑨','⑩',
           '⑪','⑫','⑬','⑭','⑮','⑯','⑰','⑱','⑲','⑳']

def _alpha_char(n): return chr(96 + ((n - 1) % 26) + 1)
def _circled(n):    return CIRCLED[(n - 1) % len(CIRCLED)]

def _number_by_depth(d, n):
    if d == 0: return f'{n}.'
    if d == 1: return f'{n})'
    if d == 2: return _circled(n)
    return f'{_alpha_char(n)})'

def _alpha_by_depth(d, n):
    if d == 0: return f'{_alpha_char(n)}.'
    if d == 1: return f'{_alpha_char(n)})'
    if d == 2: return _circled(n)
    return f'{n})'


def _render_bullet(doc, content, base_indent):
    items = content.get('items') or []
    marker = content.get('marker') or 'bullet'
    counters = [0, 0, 0, 0]
    for it in items:
        if isinstance(it, str):
            text, indent = it, 0
        else:
            text = it.get('text', '')
            indent = max(0, min(3, it.get('indent', 0)))
        if not text.strip():
            continue
        counters[indent] += 1
        for k in range(indent + 1, 4):
            counters[k] = 0
        n = counters[indent]
        if marker == 'dash':      mk = '–'
        elif marker == 'number':  mk = _number_by_depth(indent, n)
        elif marker == 'alpha':   mk = _alpha_by_depth(indent, n)
        else:                      mk = '•'

        # 항목 내 줄바꿈(\n) — 둘째 줄부터 마커 폭만큼 추가 들여쓰기
        base_left = base_indent + 0.6 * indent
        cont_left = base_left + 0.6
        lines = text.split('\n')
        total = len(lines)
        for li, ln in enumerate(lines):
            is_last = (li == total - 1)
            after_pt = 2 if is_last else 0
            if li == 0:
                _add_paragraph(doc, f'{mk}  {ln}', size=10.5,
                               indent_left=base_left, after=after_pt)
            else:
                _add_paragraph(doc, ln, size=10.5,
                               indent_left=cont_left, after=after_pt)


def _render_table(doc, content, base_indent):
    headers = content.get('headers') or []
    rows = content.get('rows') or []
    if not headers and not rows:
        return
    n_cols = max(len(headers), max((len(r) for r in rows), default=0))
    if n_cols == 0:
        return
    headers = (list(headers) + [''] * n_cols)[:n_cols]
    rows = [(list(r) + [''] * n_cols)[:n_cols] for r in rows]

    col_widths_px = content.get('col_widths') or []
    total_cm = 16.0
    if (not col_widths_px or len(col_widths_px) != n_cols
            or sum(w for w in col_widths_px if w and w > 0) <= 0):
        col_cm = [total_cm / n_cols] * n_cols
    else:
        valid = [w for w in col_widths_px if w and w > 0]
        if len(valid) < n_cols:
            avg = sum(valid) / len(valid)
            col_widths_px = [w if (w and w > 0) else avg for w in col_widths_px]
        total = sum(col_widths_px)
        col_cm = [total_cm * (w / total) for w in col_widths_px]

    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for ci, h in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        _set_cell_shading(cell, 'D9E2EC')
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        _set_font(r, size=10, bold=True)

    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            _set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lines = str(val or '').split('\n')
            for li, ln in enumerate(lines):
                if li > 0:
                    p = cell.add_paragraph()
                r = p.add_run(ln)
                _set_font(r, size=10)

    _set_table_fixed_layout(tbl, total_cm, col_cm)
    # 헤더 행 반복 (페이지 넘어가면 자동으로 위에 다시 표시)
    _set_row_as_header(tbl.rows[0])
    _add_paragraph(doc, '', before=2, after=4)


def _set_row_as_header(row):
    """이 행을 '제목 행 반복' 행으로 설정 — 표가 페이지를 넘어가면 자동 반복."""
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    if trPr.find(qn('w:tblHeader')) is None:
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), 'true')
        trPr.append(tblHeader)
    if trPr.find(qn('w:cantSplit')) is None:
        trPr.append(OxmlElement('w:cantSplit'))


def _render_image(doc, content, base_indent):
    images = content.get('images') or []
    columns = max(1, min(4, int(content.get('columns', 2) or 2)))
    if not images:
        return
    n_rows = (len(images) + columns - 1) // columns
    total_cm = 16.0
    cell_cm = (total_cm - 0.3 * (columns - 1)) / columns
    img_w = cell_cm - 0.4
    img_h = img_w * 3 / 4

    tbl = doc.add_table(rows=n_rows * 2, cols=columns)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, img in enumerate(images):
        ri = (idx // columns) * 2
        ci = idx % columns
        img_cell = tbl.rows[ri].cells[ci]
        cap_cell = tbl.rows[ri + 1].cells[ci]

        img_path = _resolve_brep_image_path(img.get('url') or '',
                                             img.get('filename') or '')
        if img_path and os.path.exists(img_path):
            try:
                processed = _crop_to_aspect(img_path, 4/3)
                if processed != img_path:
                    _GLOBAL_TEMP_FILES.append(processed)
                p = img_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(processed, width=Cm(img_w), height=Cm(img_h))
            except Exception as e:
                p = img_cell.paragraphs[0]
                rr = p.add_run(f'[이미지 로드 실패: {e}]')
                _set_font(rr, size=9, color='B91C1C')
        else:
            p = img_cell.paragraphs[0]
            rr = p.add_run('[이미지 없음]')
            _set_font(rr, size=9, color='9CA3AF')

        cap_p = cap_cell.paragraphs[0]
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = img.get('caption') or ''
        if caption:
            cr = cap_p.add_run(caption)
            _set_font(cr, size=9, color='4B5563')
            cr.italic = True

    col_cm_list = [cell_cm] * columns
    _set_table_fixed_layout(tbl, total_cm, col_cm_list)
    _add_paragraph(doc, '', before=2, after=6)


def _contain_in_static(candidate):
    """candidate 경로가 static 디렉터리 밖(경로순회 ../)이면 None 반환."""
    from app import app
    root = os.path.realpath(app.static_folder)
    real = os.path.realpath(candidate)
    if real == root or real.startswith(root + os.sep):
        return real
    return None


def _resolve_brep_image_path(url, filename):
    """boarding/ 폴더에 저장된 이미지 경로 해결 (static 디렉터리 밖 접근 차단)"""
    if url and url.startswith('/static/'):
        rel = url[len('/static/'):]
        from app import app
        return _contain_in_static(os.path.join(app.static_folder, rel))
    if filename:
        from app import app
        return _contain_in_static(
            os.path.join(app.static_folder, 'uploads', 'boarding', filename))
    return None


# ─────────────────────────────────────────────────────────────
#  신규 블록: info_table (Label-Value 표)
# ─────────────────────────────────────────────────────────────
def _render_info_table(doc, content, base_indent):
    rows = content.get('rows') or []
    rows = [r for r in rows if (r.get('label') or '').strip() or (r.get('value') or '').strip()]
    if not rows:
        return

    total_cm = 16.0
    col_cm = [4.5, total_cm - 4.5]

    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.autofit = False
    for ri, r in enumerate(rows):
        # Label 셀
        lcell = tbl.rows[ri].cells[0]
        _set_cell_shading(lcell, 'F3F4F6')
        _set_cell_borders(lcell)
        lcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = lcell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lr = lp.add_run(r.get('label') or '')
        _set_font(lr, size=10.5, bold=True)

        # Value 셀
        vcell = tbl.rows[ri].cells[1]
        _set_cell_borders(vcell)
        vcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vp = vcell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 줄바꿈 보존
        lines = (r.get('value') or '').split('\n')
        for li, ln in enumerate(lines):
            if li > 0:
                vp = vcell.add_paragraph()
            vr = vp.add_run(ln)
            _set_font(vr, size=10.5)

    _set_table_fixed_layout(tbl, total_cm, col_cm)
    _add_paragraph(doc, '', before=2, after=4)


# ─────────────────────────────────────────────────────────────
#  신규 블록: defect_table (Defect List)
# ─────────────────────────────────────────────────────────────
RISK_COLORS = {
    'L': {'bg': 'D1FAE5', 'fg': '065F46'},   # 초록
    'M': {'bg': 'FEF3C7', 'fg': '92400E'},   # 노랑
    'H': {'bg': 'FEE2E2', 'fg': '991B1B'},   # 빨강
}

def _render_defect_table(doc, content, base_indent):
    items = content.get('items') or []
    if not items:
        return

    # 헤더 + 데이터 행
    total_cm = 16.0
    # 컬럼: No(0.8) / Photo(4) / Description(5.6) / Rectification(5.6)
    col_cm = [0.8, 4.0, 5.6, 5.6]
    n_cols = 4

    tbl = doc.add_table(rows=1 + len(items), cols=n_cols)
    tbl.autofit = False

    # 헤더
    headers = ['', 'Item (Photo)', 'Description (Findings)', 'Rectification']
    for ci, h in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        _set_cell_shading(cell, '1F4E79')
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        _set_font(r, size=10.5, bold=True, color='FFFFFF')

    # 데이터 행
    for idx, it in enumerate(items, start=1):
        risk = (it.get('risk') or 'L').upper()
        if risk not in ('L', 'M', 'H'):
            risk = 'L'
        row_bg = RISK_COLORS[risk]['bg']

        cells = tbl.rows[idx].cells

        # No 셀
        no_cell = cells[0]
        _set_cell_shading(no_cell, row_bg)
        _set_cell_borders(no_cell)
        no_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        np_ = no_cell.paragraphs[0]
        np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = np_.add_run(str(idx))
        _set_font(nr, size=11, bold=True)

        # Photo 셀
        ph_cell = cells[1]
        _set_cell_shading(ph_cell, row_bg)
        _set_cell_borders(ph_cell)
        ph_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _render_defect_photos(ph_cell, it.get('images') or [], col_cm[1])

        # Description (Item + (Risk) + 줄바꿈 desc)
        desc_cell = cells[2]
        _set_cell_shading(desc_cell, row_bg)
        _set_cell_borders(desc_cell)
        desc_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # 첫 줄: Item 굵게 + (Risk)
        item_text = (it.get('item') or '').strip()
        first_p = desc_cell.paragraphs[0]
        first_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if item_text:
            run = first_p.add_run(item_text)
            _set_font(run, size=10.5, bold=True)
            sp = first_p.add_run(f'  ({risk})')
            _set_font(sp, size=10.5, bold=True, color=RISK_COLORS[risk]['fg'])
        else:
            sp = first_p.add_run(f'({risk})')
            _set_font(sp, size=10.5, bold=True, color=RISK_COLORS[risk]['fg'])
        # 나머지 줄: desc
        desc_text = (it.get('desc') or '').strip()
        if desc_text:
            for line in desc_text.split('\n'):
                lp = desc_cell.add_paragraph()
                lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                lr = lp.add_run(line)
                _set_font(lr, size=10)

        # Rectification
        fix_cell = cells[3]
        _set_cell_shading(fix_cell, row_bg)
        _set_cell_borders(fix_cell)
        fix_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        fix_text = (it.get('fix') or '').strip()
        fp = fix_cell.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lines = fix_text.split('\n') if fix_text else ['']
        for li, line in enumerate(lines):
            if li > 0:
                fp = fix_cell.add_paragraph()
                fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fr = fp.add_run(line)
            _set_font(fr, size=10)

    _set_table_fixed_layout(tbl, total_cm, col_cm)
    # 헤더 행 반복 (페이지 넘어가면 자동으로 위에 다시 표시)
    _set_row_as_header(tbl.rows[0])
    _add_paragraph(doc, '', before=2, after=2)

    # Risk Legend
    legend = doc.add_paragraph()
    legend.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = legend.add_run('Level of Risk:    ')
    _set_font(lr, size=10, bold=True)

    for v, label, color in [('L', 'L : Low', '065F46'),
                             ('M', 'M : Medium', '92400E'),
                             ('H', 'H : High', '991B1B')]:
        rr = legend.add_run(label)
        _set_font(rr, size=10, bold=True, color=color)
        sep = legend.add_run('     ')
        _set_font(sep, size=10)

    _add_paragraph(doc, '', before=2, after=4)


def _render_defect_photos(cell, images, cell_cm):
    """defect 항목의 사진 — 셀 안에 메인 1장 + 작은 thumb 그리드"""
    if not images:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('—')
        _set_font(r, size=10, color='9CA3AF')
        return

    main_img = images[0]
    extras = images[1:]
    img_w = cell_cm - 0.4

    # 메인 이미지
    main_path = _resolve_brep_image_path(main_img.get('url') or '',
                                          main_img.get('filename') or '')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if main_path and os.path.exists(main_path):
        try:
            processed = _crop_to_aspect(main_path, 4/3)
            if processed != main_path:
                _GLOBAL_TEMP_FILES.append(processed)
            run = p.add_run()
            run.add_picture(processed,
                            width=Cm(img_w),
                            height=Cm(img_w * 3 / 4))
        except Exception:
            r = p.add_run('[이미지 오류]')
            _set_font(r, size=8, color='B91C1C')
    else:
        r = p.add_run('[이미지 없음]')
        _set_font(r, size=8, color='9CA3AF')

    # 추가 이미지들
    if extras:
        extra_p = cell.add_paragraph()
        extra_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        thumb_w = (img_w - 0.2 * min(3, len(extras))) / min(4, len(extras))
        thumb_h = thumb_w * 3 / 4
        for ex in extras[:4]:
            ex_path = _resolve_brep_image_path(ex.get('url') or '',
                                                ex.get('filename') or '')
            if ex_path and os.path.exists(ex_path):
                try:
                    processed = _crop_to_aspect(ex_path, 4/3)
                    if processed != ex_path:
                        _GLOBAL_TEMP_FILES.append(processed)
                    run = extra_p.add_run()
                    run.add_picture(processed,
                                    width=Cm(thumb_w), height=Cm(thumb_h))
                    extra_p.add_run(' ')
                except Exception:
                    pass


# ─────────────────────────────────────────────────────────────
#  공개 함수 — 보고서 데이터 받아 docx 바이트 반환
# ─────────────────────────────────────────────────────────────
def build_docx(report: dict) -> bytes:
    doc = Document()

    # 페이지 설정
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width  = Mm(210)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    rFonts.set(qn('w:ascii'),    'Malgun Gothic')
    rFonts.set(qn('w:hAnsi'),    'Malgun Gothic')

    _add_sinokor_footer(doc)

    # 섹션 트리 빌드
    sections_flat = report.get('sections') or []
    tree = _build_tree(sections_flat)

    # ① 표지 (회사 양식: 통합 표 안에 헤더 + 정보 + 결재란)
    _build_brep_cover(doc, report)

    # ② 본문 — 1단계 섹션을 회사 양식의 큰 통합 표(좌측 라벨/우측 본문)로 렌더링
    # defect_table은 통합 표 바깥에 별도로 (사진 + 색깔 행 필요해서)
    _render_brep_main_body(doc, tree)

    # 직렬화
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    result = bio.read()

    # 임시 파일 정리
    global _GLOBAL_TEMP_FILES
    for fp in _GLOBAL_TEMP_FILES:
        try: os.remove(fp)
        except Exception: pass
    _GLOBAL_TEMP_FILES.clear()

    return result


def _render_brep_main_body(doc, tree):
    """
    회사 양식 + 사용자 요청 구조:
      [표지 (Vessel Boarding Report 통합 표)]
      [Items | Content 헤더 행]                ← 양식의 R7
      [Inspector Opinion | 본문 (양식의 표2 row1)]
      [Vessel General Condition & Deficiencies | 본문]
      ── 페이지 나누기 ──
      [별첨 | 사진 (사진 갤러리)]
      ── 페이지 나누기 ──
      [Defect List (별도 표 + 색깔)]
    """
    if not tree:
        return

    # 본문 영역: 좌측 라벨 / 우측 본문 (양식의 표 2)
    # 컬럼: 2.42cm (라벨) + 15.26cm (본문) = 17.68cm (표지와 동일 폭)
    body_col_cm = [2.42, 15.26]
    body_total = sum(body_col_cm)

    # ─── "Items / Content" 헤더 행 (양식의 R7) ───
    items_tbl = doc.add_table(rows=1, cols=2)
    items_tbl.autofit = False
    items_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    ic = items_tbl.rows[0].cells[0]
    _set_cell_shading(ic, 'F2F2F2'); _set_cell_borders(ic)
    ic.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    ip = ic.paragraphs[0]; ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = ip.add_run('Items'); _set_font(ir, size=10, bold=True)
    cc = items_tbl.rows[0].cells[1]
    _set_cell_shading(cc, 'F2F2F2'); _set_cell_borders(cc)
    cc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cp = cc.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run('Content'); _set_font(cr, size=10, bold=True)
    _set_row_height(items_tbl.rows[0], 0.55)
    _set_table_fixed_layout(items_tbl, body_total, body_col_cm)

    # ─── 섹션 자동 매핑 ───
    # 자동 생성된 4개 섹션을 찾아서 양식 위치에 배치:
    #  1) Inspector Opinion → 표지와 같은 페이지 (Items/Content 다음)
    #  2) Vessel General Condition → Inspector Opinion 다음
    #  3) 첨부 사진 → 페이지 나누기 후
    #  4) Defect List → 또 페이지 나누기 후
    import re

    sec_by_kind = {'opinion': None, 'condition': None, 'photo': None, 'defect': None}
    other_secs = []

    def classify(title):
        t = (title or '').lower()
        if 'inspector' in t or 'opinion' in t or '의견' in t:
            return 'opinion'
        if 'vessel' in t or 'condition' in t or 'deficien' in t or '컨디션' in t:
            return 'condition'
        if '사진' in t or 'photo' in t or '첨부' in t or '별첨' in t:
            return 'photo'
        if 'defect' in t or '결함' in t or '디펙' in t:
            return 'defect'
        return None

    for sec in tree:
        kind = classify(sec.get('title', ''))
        if kind and not sec_by_kind[kind]:
            sec_by_kind[kind] = sec
        else:
            other_secs.append(sec)

    def clean_label(s):
        return re.sub(r'^\s*\d+[.)]\s*', '', (s or '').strip())

    # ─── 1) Inspector Opinion 행 ───
    if sec_by_kind['opinion']:
        _render_body_row(doc, sec_by_kind['opinion'],
                         label_override='Inspector\nopinion',
                         col_cm=body_col_cm, total_cm=body_total)
    else:
        # 빈 행
        _render_body_row_empty(doc, 'Inspector\nopinion', body_col_cm, body_total)

    # ─── 2) Vessel General Condition & Deficiencies 행 ───
    if sec_by_kind['condition']:
        _render_body_row(doc, sec_by_kind['condition'],
                         label_override='Vessel general\ncondition &\nDeficiencies',
                         col_cm=body_col_cm, total_cm=body_total)
    else:
        _render_body_row_empty(doc, 'Vessel general\ncondition &\nDeficiencies',
                               body_col_cm, body_total)

    # 다른 (자동 분류 안 된) 섹션이 있으면 여기에 같은 양식으로 추가
    for sec in other_secs:
        title = clean_label(sec.get('title'))
        _render_body_row(doc, sec, label_override=title,
                         col_cm=body_col_cm, total_cm=body_total)

    # ─── 페이지 나누기 → 첨부 사진 (Defect List와 동일 스타일: 제목 + 콘텐츠) ───
    if sec_by_kind['photo']:
        doc.add_page_break()
        _add_paragraph(doc, '첨부 사진',
                       size=14, bold=True, color='1F4E79',
                       before=4, after=8)
        # 사진 블록만 외곽 표 없이 본문 영역 전체 폭에 출력
        blocks = sorted(sec_by_kind['photo'].get('blocks') or [],
                        key=lambda b: (b.get('display_order', 0), b.get('id', 0)))
        for b in blocks:
            content = b.get('content') or {}
            if isinstance(content, str):
                try: content = json.loads(content)
                except Exception: content = {}
            if b.get('block_type') == 'image':
                # 본문 폭 17.68cm 가득 사용
                _render_image_full_width(doc, content, total_cm=17.68)
            else:
                _render_brep_block(doc, b, 0)

    # ─── 페이지 나누기 → Defect List (별도 표 + 색깔) ───
    if sec_by_kind['defect']:
        doc.add_page_break()
        # 라벨 행
        _add_paragraph(doc, 'Defect List',
                       size=14, bold=True, color='1F4E79',
                       before=4, after=6)
        # defect_table 블록들 별도 출력
        blocks = sorted(sec_by_kind['defect'].get('blocks') or [],
                        key=lambda b: (b.get('display_order', 0), b.get('id', 0)))
        had_defect = False
        for b in blocks:
            content = b.get('content') or {}
            if isinstance(content, str):
                try: content = json.loads(content)
                except Exception: content = {}
            if b.get('block_type') == 'defect_table':
                _render_defect_table(doc, content, 0)
                had_defect = True
            else:
                # 다른 블록 타입은 일반 렌더링
                _render_brep_block(doc, b, 0)
        if not had_defect:
            _add_paragraph(doc, '(아직 작성된 결함 항목이 없습니다.)',
                           size=10, color='6B7280', after=4)


def _render_image_full_width(doc, content, total_cm=17.68):
    """좌측 라벨 셀 없이 본문 폭 가득 사용하는 이미지 갤러리"""
    images = content.get('images') or []
    columns = max(1, min(4, int(content.get('columns', 2) or 2)))
    if not images:
        return

    n_rows = (len(images) + columns - 1) // columns
    cell_cm = total_cm / columns
    img_w = cell_cm - 0.4
    img_h = img_w * 3 / 4

    tbl = doc.add_table(rows=n_rows * 2, cols=columns)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, img in enumerate(images):
        ri = (idx // columns) * 2
        ci = idx % columns
        img_cell = tbl.rows[ri].cells[ci]
        cap_cell = tbl.rows[ri + 1].cells[ci]
        _set_cell_borders(img_cell, color='D1D5DB')
        _set_cell_borders(cap_cell, color='D1D5DB')

        img_path = _resolve_brep_image_path(img.get('url') or '',
                                             img.get('filename') or '')
        p = img_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if img_path and os.path.exists(img_path):
            try:
                processed = _crop_to_aspect(img_path, 4/3)
                if processed != img_path:
                    _GLOBAL_TEMP_FILES.append(processed)
                run = p.add_run()
                run.add_picture(processed, width=Cm(img_w), height=Cm(img_h))
            except Exception:
                r = p.add_run('[이미지 오류]')
                _set_font(r, size=9, color='B91C1C')
        else:
            r = p.add_run('[이미지 없음]')
            _set_font(r, size=9, color='9CA3AF')

        cap_p = cap_cell.paragraphs[0]
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = img.get('caption') or ''
        if caption:
            cr = cap_p.add_run(caption)
            _set_font(cr, size=9, color='4B5563')
            cr.italic = True

    col_cm_list = [cell_cm] * columns
    _set_table_fixed_layout(tbl, total_cm, col_cm_list)
    _add_paragraph(doc, '', before=2, after=4)


def _render_body_row(doc, sec, label_override, col_cm, total_cm):
    """양식의 표2 row 1개: 좌측 라벨 / 우측 본문 (블록들)"""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 좌측 라벨 셀
    lc = tbl.rows[0].cells[0]
    _set_cell_shading(lc, 'F2F2F2'); _set_cell_borders(lc)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = lc.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run(label_override)
    _set_font(lr, size=10, bold=True)

    # 우측 본문 셀
    rc = tbl.rows[0].cells[1]
    _set_cell_borders(rc)
    rc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    # 기본 paragraph 제거
    p0 = rc.paragraphs[0]
    p0._element.getparent().remove(p0._element)

    # 모든 블록을 우측 셀에 (defect_table은 별도라 제외)
    blocks = sorted(sec.get('blocks') or [],
                    key=lambda b: (b.get('display_order', 0), b.get('id', 0)))
    rendered = False
    for b in blocks:
        if b.get('block_type') == 'defect_table':
            continue
        _render_brep_block_in_cell(rc, b)
        rendered = True
    if not rendered:
        # 빈 셀에 빈 paragraph 추가 (테두리 유지를 위해)
        rc.add_paragraph()

    _set_table_fixed_layout(tbl, total_cm, col_cm)


def _render_body_row_empty(doc, label, col_cm, total_cm):
    """비어있는 행 (라벨만)"""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    lc = tbl.rows[0].cells[0]
    _set_cell_shading(lc, 'F2F2F2'); _set_cell_borders(lc)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = lc.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run(label); _set_font(lr, size=10, bold=True)

    rc = tbl.rows[0].cells[1]
    _set_cell_borders(rc)
    _set_row_height(tbl.rows[0], 1.5)
    _set_table_fixed_layout(tbl, total_cm, col_cm)


def _build_tree(sections_flat):
    by_id = {s['id']: dict(s, children=[]) for s in sections_flat}
    roots = []
    for s in by_id.values():
        if s.get('parent_id') and s['parent_id'] in by_id:
            by_id[s['parent_id']]['children'].append(s)
        else:
            roots.append(s)
    def sort_rec(lst):
        lst.sort(key=lambda x: (x.get('display_order', 0), x.get('id', 0)))
        for x in lst:
            sort_rec(x['children'])
    sort_rec(roots)
    return roots


def _add_sinokor_footer(doc):
    """첨부 양식의 푸터:
       'CODE<107-301>/2015.04.17     Sinokor Ship Management Co., Ltd'
       좌측: CODE / 우측: Sinokor — 탭으로 양쪽 정렬
    """
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]

        # 탭 스톱: 우측 정렬 위치 (본문 가용폭 16cm 기준)
        pPr = p._element.get_or_add_pPr()
        existing_tabs = pPr.find(qn('w:tabs'))
        if existing_tabs is not None:
            pPr.remove(existing_tabs)
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '9355')   # ~16.5cm in twips
        tabs.append(tab)
        pPr.append(tabs)

        r1 = p.add_run('CODE<107-301>/2015.04.17')
        _set_font(r1, size=8, color='6B7280')

        p.add_run('\t')

        r2 = p.add_run('Sinokor Ship Management Co., Ltd')
        _set_font(r2, size=8, color='6B7280', bold=True)
