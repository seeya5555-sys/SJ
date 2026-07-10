"""
Dry Dock Report — Word(.docx) 생성

첨부 양식 (PACIFIC BUSAN호) 구조를 그대로 재현:
  · 표지: 보고서 제목 + 선박 정보 표 + 결재란
  · 본문: 1단계 큰 제목 → 1) 2) 3) 하위 → 본문 (불릿/표/사진)
  · 1단계와 1단계 사이는 빈 줄로만 구분 (페이지 분리 X)
  · 사용된 마커: 1단계="1. 2.", 2단계="1) 2)", 불릿마커는 원본 그대로(•, –, 1), a))
"""
import io
import json
import os
from typing import Dict, List, Tuple
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────
#  XML 유틸 (셀 음영 / 페이지 설정 등)
# ─────────────────────────────────────────────────────────────
def _set_cell_shading(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)


def _set_cell_borders(cell, color='808080', sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for side in ['top', 'left', 'bottom', 'right']:
        b = tcBorders.find(qn(f'w:{side}'))
        if b is None:
            b = OxmlElement(f'w:{side}')
            tcBorders.append(b)
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:color'), color)


def _set_font(run, *, name='Malgun Gothic', size=10, bold=False, color=None):
    run.font.name = name
    # 한글 폰트도 함께 지정
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color.lstrip('#'))


def _add_paragraph(doc_or_cell, text='', *, font='Malgun Gothic', size=10, bold=False,
                   color=None, align=None, before=0, after=0, indent_left=0):
    p = doc_or_cell.add_paragraph()
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if indent_left:
        pf.left_indent = Cm(indent_left)
    if text:
        r = p.add_run(text)
        _set_font(r, name=font, size=size, bold=bold, color=color)
    return p


# ─────────────────────────────────────────────────────────────
#  표지 — 결재란 + 보고서 제목 + 선박 정보 표
# ─────────────────────────────────────────────────────────────
def _build_cover(doc, report):
    """첨부 양식의 표지 페이지"""

    # ① 결재란 (우측 정렬) — 단일 표
    approvals = [
        ('기 안', report.get('approval_drafter') or ''),
        ('팀 장', report.get('approval_team_lead') or ''),
        ('중 역', report.get('approval_director') or ''),
        ('대표이사', report.get('approval_ceo') or ''),
    ]

    # 결재 라벨 1열 + 결재자 4열 = 총 5열, 2행
    approval_tbl = doc.add_table(rows=2, cols=len(approvals) + 1)
    approval_tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
    approval_tbl.autofit = False

    # 표 자체에 left indent를 줘서 우측으로 밀어붙임
    # 본문폭 17cm, 표 총 9.4cm → 좌측 들여쓰기 ~7.5cm
    tblPr = approval_tbl._element.find(qn('w:tblPr'))
    if tblPr is not None:
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '4250')  # twips (약 7.5cm)
        tblInd.set(qn('w:type'), 'dxa')
        # tblPr 안에서 적절한 위치에 삽입
        existing = tblPr.find(qn('w:tblInd'))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(tblInd)

    col_widths_cm = [1.4] + [2.0] * len(approvals)
    # 첫 열: "결재" 라벨 (세로 병합)
    title_cell = approval_tbl.rows[0].cells[0]
    title_cell = title_cell.merge(approval_tbl.rows[1].cells[0])
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    title_cell.width = Cm(col_widths_cm[0])
    _set_cell_shading(title_cell, 'F2F2F2')
    # "결재" 한 줄로 (Word에서 잘 보이도록)
    p = title_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('결재')
    _set_font(r, size=11, bold=True)
    # 4면 모두 명시적으로 테두리 설정 (병합 셀에서 누락 방지)
    _set_cell_borders_sides(title_cell, {'top', 'left', 'bottom', 'right'},
                            color='808080', sz=4)
    # 병합된 결재 셀의 너비도 XML 레벨에서 강제 (LibreOffice가 무시하는 경우 방지)
    _tcPr = title_cell._tc.get_or_add_tcPr()
    _existing_tcw = _tcPr.find(qn('w:tcW'))
    if _existing_tcw is not None:
        _tcPr.remove(_existing_tcw)
    _tcW = OxmlElement('w:tcW')
    _tcW.set(qn('w:w'), str(int(col_widths_cm[0] * 567)))
    _tcW.set(qn('w:type'), 'dxa')
    _tcPr.append(_tcW)

    # 결재자 헤더 + 빈 서명란
    for col, (label, name) in enumerate(approvals, start=1):
        h = approval_tbl.rows[0].cells[col]
        h.width = Cm(col_widths_cm[col])
        _set_cell_shading(h, 'F2F2F2')
        h.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        ph = h.paragraphs[0]
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rh = ph.add_run(label)
        _set_font(rh, size=9, bold=True)
        _set_cell_borders(h)

        sig = approval_tbl.rows[1].cells[col]
        sig.width = Cm(col_widths_cm[col])
        sig.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        ps = sig.paragraphs[0]
        ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if name:
            rs = ps.add_run(name)
            _set_font(rs, size=10)
        _set_cell_borders(sig)

    # 서명란 행 높이 고정
    _set_row_height(approval_tbl.rows[1], 1.3)

    # 표 전체 너비 고정 (셀 너비 확정)
    _set_table_fixed_layout(approval_tbl,
                            sum(col_widths_cm), col_widths_cm)

    _add_paragraph(doc, '', before=12)

    # ② 보고서 제목
    title = report.get('title') or 'Dry Dock Report'
    _add_paragraph(doc, title,
                   size=18, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   before=24, after=8)

    # 제목 밑줄선
    line_p = doc.add_paragraph()
    pPr = line_p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)

    _add_paragraph(doc, '', before=8)

    # ③ 선박 정보 표 (라벨 | 값 | 라벨 | 값) - 2행 4열
    info_rows = [
        ('Vessel Name', report.get('vessel_name') or '',
         'Type', report.get('vessel_type') or ''),
        ('IMO No.',    report.get('imo_no') or '',
         'Built',      ''),  # built 정보는 양식에 없으나 자리 유지
        ('Gross Tonnage', report.get('gross_tonnage') or '',
         'Dead Weight',   report.get('dead_weight') or ''),
        ('Shipyard',  report.get('shipyard') or '',
         'Dock No.',  report.get('dock_no') or ''),
        ('Dry Dock Period',
         _fmt_period(report.get('period_start'), report.get('period_end')),
         'Reported on', datetime.now().strftime('%Y-%m-%d')),
    ]

    info_tbl = doc.add_table(rows=len(info_rows), cols=4)
    info_tbl.autofit = False
    widths = [Cm(3.5), Cm(5.5), Cm(3.5), Cm(5.5)]
    for ri, row in enumerate(info_rows):
        for ci, val in enumerate(row):
            cell = info_tbl.rows[ri].cells[ci]
            cell.width = widths[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_borders(cell)
            p = cell.paragraphs[0]
            # 라벨 셀(짝수 인덱스): 음영 + 굵게
            if ci % 2 == 0:
                _set_cell_shading(cell, 'E7EBF0')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(val)
                _set_font(r, size=10, bold=True)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(val)
                _set_font(r, size=10)

    # 페이지 나누기 (본문은 새 페이지부터)
    doc.add_page_break()


def _set_no_border(cell, side):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    b = tcBorders.find(qn(f'w:{side}'))
    if b is None:
        b = OxmlElement(f'w:{side}')
        tcBorders.append(b)
    b.set(qn('w:val'), 'nil')


def _set_row_height(row, cm):
    """행 높이 강제 설정 (cm)"""
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(cm * 567)))  # 1cm = 567 twips
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def _fmt_period(start, end):
    if not start and not end:
        return ''
    s = (start or '').replace('-', '.')
    e = (end or '').replace('-', '.')
    if s and e:
        try:
            d1 = datetime.strptime(start, '%Y-%m-%d')
            d2 = datetime.strptime(end, '%Y-%m-%d')
            days = (d2 - d1).days + 1
            return f'{s} ~ {e}  ({days}일)'
        except Exception:
            return f'{s} ~ {e}'
    return s or e


# ─────────────────────────────────────────────────────────────
#  목차 (Table of Contents) — 자동 생성, 페이지 번호 없이 트리 형태
# ─────────────────────────────────────────────────────────────
def _build_toc(doc, sections_tree):
    _add_paragraph(doc, '목 차',
                   size=16, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   before=12, after=18)

    # tree 평면화 + 번호링
    def walk(nodes, prefix='', depth=0):
        for i, n in enumerate(nodes):
            num = f'{prefix}-{i + 1}' if prefix else f'{i + 1}'
            yield depth, num, n
            yield from walk(n.get('children', []), num, depth + 1)

    for depth, num, n in walk(sections_tree):
        if depth == 0:
            text = f'{num}.  {n["title"]}'
            _add_paragraph(doc, text, size=11.5, bold=True,
                           before=4, after=2)
        elif depth == 1:
            text = f'{num}.  {n["title"]}'   # "1-1." 형식
            _add_paragraph(doc, text, size=10.5, indent_left=0.8,
                           before=2, after=1)
        else:
            text = f'{num}.  {n["title"]}'   # "1-1-1." 형식
            _add_paragraph(doc, text, size=10, indent_left=1.6,
                           before=1, after=1)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
#  본문 — 섹션 트리 + 블록 렌더링
# ─────────────────────────────────────────────────────────────
def _render_sections(doc, sections_tree, depth=0, prefix=''):
    """재귀적으로 섹션과 그 아래 블록 + 자식 섹션 렌더링"""
    for i, sec in enumerate(sections_tree):
        num = f'{prefix}-{i + 1}' if prefix else f'{i + 1}'

        # 1단계가 두 번째 이상이면 새 페이지로 시작
        if depth == 0 and i > 0:
            doc.add_page_break()

        if depth == 0:
            # 1단계 — 큰 제목 ("1.")
            _add_paragraph(doc, f'{num}. {sec["title"]}',
                           size=14, bold=True, color='1F4E79',
                           before=18, after=8)
            # 제목 밑 선
            _add_horizontal_line(doc, color='1F4E79', sz=8)
        elif depth == 1:
            # 2단계 — "1-1. 제목" 형식 (1단계와 같은 스타일이지만 글자 크기만 약간 작음)
            _add_paragraph(doc, f'{num}. {sec["title"]}',
                           size=13, bold=True, color='1F4E79',
                           before=14, after=6)
            _add_horizontal_line(doc, color='5B9BD5', sz=6)
        else:
            # 3단계 — "1-1-1. 제목" 형식, 더 작은 제목
            _add_paragraph(doc, f'{num}. {sec["title"]}',
                           size=11.5, bold=True, color='2E5990',
                           before=10, after=4)

        # 블록 렌더링
        blocks = sorted(sec.get('blocks', []),
                        key=lambda b: (b.get('display_order', 0), b.get('id', 0)))
        for b in blocks:
            _render_block(doc, b, depth)

        # 자식 섹션 (같은 페이지에 이어서)
        if sec.get('children'):
            _render_sections(doc, sec['children'], depth + 1, num)


def _add_horizontal_line(doc, color='808080', sz=6):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _render_block(doc, block, depth):
    """블록 타입별 렌더링"""
    bt = block.get('block_type')
    content = block.get('content') or {}
    # content가 문자열로 저장된 경우 (DB raw row) 디시리얼라이즈
    if isinstance(content, str):
        try:
            content = json.loads(content)
        except Exception:
            content = {}

    # depth에 따른 좌측 들여쓰기 (1단계 본문은 0.5cm, 2단계는 1cm)
    base_indent = 0.5 + 0.5 * depth

    if bt == 'paragraph':
        text = (content.get('text') or '').strip()
        if not text:
            return
        for line in text.split('\n'):
            _add_paragraph(doc, line, size=10.5,
                           indent_left=base_indent, after=4)

    elif bt == 'bullet_list':
        items = content.get('items') or []
        marker = content.get('marker') or 'bullet'
        # 레벨별 카운터
        counters = [0, 0, 0, 0]
        # 원숫자 (①~⑳)
        CIRCLED = ['①','②','③','④','⑤','⑥','⑦','⑧','⑨','⑩',
                   '⑪','⑫','⑬','⑭','⑮','⑯','⑰','⑱','⑲','⑳']

        def _alpha(n):
            return chr(96 + ((n - 1) % 26) + 1)

        def _circled(n):
            return CIRCLED[(n - 1) % len(CIRCLED)]

        def _number_by_depth(depth, n):
            if depth == 0: return f'{n}.'
            if depth == 1: return f'{n})'
            if depth == 2: return _circled(n)
            return f'{_alpha(n)})'

        def _alpha_by_depth(depth, n):
            if depth == 0: return f'{_alpha(n)}.'
            if depth == 1: return f'{_alpha(n)})'
            if depth == 2: return _circled(n)
            return f'{n})'

        for it in items:
            if isinstance(it, str):
                text, indent = it, 0
            else:
                text = it.get('text', '')
                indent = max(0, min(3, it.get('indent', 0)))
            if not text.strip():
                continue
            counters[indent] += 1
            for k in range(indent + 1, 4):
                counters[k] = 0

            n = counters[indent]
            if marker == 'dash':      mk = '–'
            elif marker == 'number':  mk = _number_by_depth(indent, n)
            elif marker == 'alpha':   mk = _alpha_by_depth(indent, n)
            else:                      mk = '•'

            # 항목 내 줄바꿈(\n) 처리:
            #   첫 줄  → 마커 + text  (들여쓰기 = base + 0.6*indent)
            #   둘째 줄~ → 마커 없이    (추가로 마커 폭만큼 들여쓰기, hanging indent)
            base_left = base_indent + 0.6 * indent
            cont_left = base_left + 0.6   # 마커("•  ") 폭만큼 추가 들여쓰기
            lines = text.split('\n')
            total = len(lines)
            for li, ln in enumerate(lines):
                is_last = (li == total - 1)
                after_pt = 2 if is_last else 0
                if li == 0:
                    _add_paragraph(doc, f'{mk}  {ln}', size=10.5,
                                   indent_left=base_left, after=after_pt)
                else:
                    _add_paragraph(doc, ln, size=10.5,
                                   indent_left=cont_left, after=after_pt)

    elif bt == 'table':
        _render_table_block(doc, content, base_indent)

    elif bt == 'image':
        _render_image_block(doc, content, base_indent)


def _set_table_fixed_layout(tbl, total_cm, col_cm_list):
    """
    표 전체 폭을 고정하고 컬럼 너비를 정확히 적용.
      · tblLayout=fixed: Word가 셀 내용에 따라 너비 자동 조정 못 하게
      · tblW: 표 전체 너비 강제
      · w:gridCol: 각 컬럼 너비 (Word가 우선 참조)
    """
    tblPr = tbl._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._element.insert(0, tblPr)

    # 1) fixed layout
    existing_layout = tblPr.find(qn('w:tblLayout'))
    if existing_layout is not None:
        tblPr.remove(existing_layout)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)

    # 2) 전체 너비
    existing_w = tblPr.find(qn('w:tblW'))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(total_cm * 567)))  # cm → twips (1cm = 567 twips)
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # 3) grid (각 컬럼 너비) — Word는 이걸 우선 참조
    existing_grid = tbl._element.find(qn('w:tblGrid'))
    if existing_grid is not None:
        tbl._element.remove(existing_grid)
    grid = OxmlElement('w:tblGrid')
    for w_cm in col_cm_list:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(int(w_cm * 567)))
        grid.append(col)
    # tblPr 다음에 삽입
    tblPr.addnext(grid)

    # 4) 셀별 너비도 명시 (안정성)
    for row in tbl.rows:
        for ci, cell in enumerate(row.cells):
            if ci >= len(col_cm_list):
                continue
            tcPr = cell._tc.get_or_add_tcPr()
            existing_tcw = tcPr.find(qn('w:tcW'))
            if existing_tcw is not None:
                tcPr.remove(existing_tcw)
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(col_cm_list[ci] * 567)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)


def _normalize_table_content(content):
    """기존 {headers, rows} 또는 새 {cells, header_row_count} 모두 받아서
       통일된 {cells, header_row_count, col_widths} 반환"""
    if isinstance(content.get('cells'), list):
        cells = content['cells']
        # 각 셀 정규화
        norm_cells = []
        for row in cells:
            norm_row = []
            for c in row:
                if c is None:
                    norm_row.append(None)
                else:
                    norm_row.append({
                        'text': c.get('text', '') or '',
                        'rowspan': max(1, int(c.get('rowspan', 1))),
                        'colspan': max(1, int(c.get('colspan', 1))),
                    })
            norm_cells.append(norm_row)
        return {
            'cells': norm_cells,
            'header_row_count': max(1, int(content.get('header_row_count', 1))),
            'col_widths': list(content.get('col_widths') or []),
        }

    # 옛 구조 → 변환
    headers = content.get('headers') or []
    rows = content.get('rows') or []
    n_cols = max(len(headers), max((len(r) for r in rows), default=0))
    if n_cols == 0:
        return None
    headers = (list(headers) + [''] * n_cols)[:n_cols]
    rows = [(list(r) + [''] * n_cols)[:n_cols] for r in rows]
    cells = [
        [{'text': h or '', 'rowspan': 1, 'colspan': 1} for h in headers],
    ]
    for r in rows:
        cells.append([{'text': v or '', 'rowspan': 1, 'colspan': 1} for v in r])
    return {
        'cells': cells,
        'header_row_count': 1,
        'col_widths': list(content.get('col_widths') or []),
    }


def _render_table_block(doc, content, base_indent):
    norm = _normalize_table_content(content)
    if not norm or not norm['cells']:
        return

    cells = norm['cells']
    header_row_count = norm['header_row_count']
    col_widths_px = norm['col_widths']

    n_rows = len(cells)
    n_cols = max((len(r) for r in cells), default=0)
    if n_cols == 0:
        return

    # 컬럼 너비 (cm)
    # A4 21cm - 좌우 마진 2cm씩 = 본문 가용 폭 17cm
    total_cm = 17.0
    if (not col_widths_px or len(col_widths_px) != n_cols
            or sum(w for w in col_widths_px if w and w > 0) <= 0):
        col_cm = [total_cm / n_cols] * n_cols
    else:
        valid = [w for w in col_widths_px if w and w > 0]
        if len(valid) < n_cols:
            avg = sum(valid) / len(valid)
            col_widths_px = [w if (w and w > 0) else avg for w in col_widths_px]
        total = sum(col_widths_px)
        col_cm = [total_cm * (w / total) for w in col_widths_px]

    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 1) 모든 마스터 셀에 텍스트 채우기
    for ri in range(n_rows):
        for ci in range(n_cols):
            cell_data = cells[ri][ci]
            if cell_data is None:
                continue
            cell = tbl.rows[ri].cells[ci]
            is_header = (ri < header_row_count)
            if is_header:
                _set_cell_shading(cell, 'D9E2EC')
            _set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # 첫 paragraph 초기화
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
            text = str(cell_data.get('text') or '')
            lines = text.split('\n')
            for li, ln in enumerate(lines):
                if li > 0:
                    p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(ln)
                _set_font(r, size=10, bold=is_header)

    # 2) 병합 적용 (rowspan/colspan > 1인 마스터 셀)
    for ri in range(n_rows):
        for ci in range(n_cols):
            cell_data = cells[ri][ci]
            if cell_data is None:
                continue
            rs = cell_data.get('rowspan', 1)
            cs = cell_data.get('colspan', 1)
            if rs > 1 or cs > 1:
                r_end = min(n_rows - 1, ri + rs - 1)
                c_end = min(n_cols - 1, ci + cs - 1)
                top_left = tbl.cell(ri, ci)
                bottom_right = tbl.cell(r_end, c_end)
                try:
                    top_left.merge(bottom_right)
                except Exception:
                    pass

    # 3) 헤더 행 반복: 페이지가 넘어가면 자동으로 맨 위에 다시 표시
    for ri in range(min(header_row_count, n_rows)):
        _set_row_as_header(tbl.rows[ri])

    # 표 너비 고정
    _set_table_fixed_layout(tbl, total_cm, col_cm)
    _add_paragraph(doc, '', before=2, after=4)


def _set_row_as_header(row):
    """이 행을 '제목 행 반복' 행으로 설정 — 표가 페이지를 넘어가면 자동 반복.
       Word의 <w:trPr><w:tblHeader/></w:trPr> 속성 추가."""
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    # 기존 tblHeader 있으면 중복 방지
    existing = trPr.find(qn('w:tblHeader'))
    if existing is None:
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), 'true')
        trPr.append(tblHeader)
    # 페이지가 분리되어도 행 자체가 잘리지 않도록 cantSplit도 추가
    cantSplit = trPr.find(qn('w:cantSplit'))
    if cantSplit is None:
        cs = OxmlElement('w:cantSplit')
        trPr.append(cs)


def _crop_to_aspect(src_path, target_ratio=4/3):
    """
    이미지를 target_ratio(가로/세로)에 맞춰 center-crop한 임시 파일 경로 반환.
    원본이 이미 비율 맞으면 그대로 원본 경로 반환.
    실패 시 원본 경로 반환.
    """
    try:
        from PIL import Image
        import tempfile
        with Image.open(src_path) as im:
            # EXIF orientation 적용
            try:
                from PIL import ImageOps
                im = ImageOps.exif_transpose(im)
            except Exception:
                pass

            w, h = im.size
            if w <= 0 or h <= 0:
                return src_path
            cur_ratio = w / h
            # 이미 비율이 비슷하면 그대로
            if abs(cur_ratio - target_ratio) < 0.02:
                return src_path

            if cur_ratio > target_ratio:
                # 너무 가로로 길다 → 좌우 잘라냄
                new_w = int(h * target_ratio)
                left = (w - new_w) // 2
                box = (left, 0, left + new_w, h)
            else:
                # 너무 세로로 길다 → 위아래 잘라냄
                new_h = int(w / target_ratio)
                top = (h - new_h) // 2
                box = (0, top, w, top + new_h)

            cropped = im.crop(box)
            # RGBA(투명PNG)는 JPEG 저장 위해 RGB로
            if cropped.mode in ('RGBA', 'P', 'LA'):
                cropped = cropped.convert('RGB')

            # 임시 파일 저장 (JPEG, quality 88)
            fd, tmp_path = tempfile.mkstemp(suffix='.jpg', prefix='dock_img_')
            os.close(fd)
            cropped.save(tmp_path, 'JPEG', quality=88, optimize=True)
            return tmp_path
    except Exception as e:
        # 실패 시 원본 사용
        return src_path


def _set_cell_borders_sides(cell, sides, color='808080', sz=4):
    """셀의 특정 면(top/left/bottom/right)만 테두리 설정. sides는 set 또는 list."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for side in ['top', 'left', 'bottom', 'right']:
        b = tcBorders.find(qn(f'w:{side}'))
        if b is None:
            b = OxmlElement(f'w:{side}')
            tcBorders.append(b)
        if side in sides:
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), str(sz))
            b.set(qn('w:color'), color)
        else:
            b.set(qn('w:val'), 'nil')


def _set_cell_vertical_padding(cell, top_twips=20, bottom_twips=20):
    """셀의 상하 안쪽 여백(margin) 설정. 기본 Word 값은 ~80~120 twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for side, val in [('top', top_twips), ('bottom', bottom_twips)]:
        e = tcMar.find(qn(f'w:{side}'))
        if e is None:
            e = OxmlElement(f'w:{side}')
            tcMar.append(e)
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')


def _render_image_block(doc, content, base_indent):
    images = content.get('images') or []
    columns = max(1, min(4, int(content.get('columns', 2) or 2)))

    if not images:
        return

    n = len(images)
    n_rows = (n + columns - 1) // columns

    # A4 21cm - 좌우 마진 2cm씩 = 본문 가용 폭 17cm
    total_cm = 17.0
    cell_cm = (total_cm - 0.3 * (columns - 1)) / columns
    img_width_cm = cell_cm - 0.4
    img_height_cm = img_width_cm * 3 / 4

    # 각 그리드 행마다 캡션 행이 필요한지 미리 결정
    # (그 행의 이미지 중 하나라도 캡션이 있으면 캡션 행 추가)
    row_has_caption = []
    for gr in range(n_rows):
        start = gr * columns
        end   = min(start + columns, n)
        has = any((images[i].get('caption') or '').strip() for i in range(start, end))
        row_has_caption.append(has)

    # 실제 표의 행 수 = 그리드행마다 1(이미지) + (캡션있으면 +1)
    total_rows = sum(1 + (1 if h else 0) for h in row_has_caption)

    tbl = doc.add_table(rows=total_rows, cols=columns)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    border_color = 'D1D5DB'
    temp_files = []

    # 그리드 행 → 실제 표의 시작 행 인덱스 매핑
    row_offsets = []
    cursor = 0
    for h in row_has_caption:
        row_offsets.append(cursor)
        cursor += 2 if h else 1

    for idx, img in enumerate(images):
        gr = idx // columns
        ci = idx % columns
        img_row_idx = row_offsets[gr]
        has_caption = row_has_caption[gr]

        img_cell = tbl.rows[img_row_idx].cells[ci]
        cap_cell = tbl.rows[img_row_idx + 1].cells[ci] if has_caption else None

        # 테두리:
        #  - 캡션 있는 경우: 이미지(top/left/right), 캡션(bottom/left/right) — 사이는 비움
        #  - 캡션 없는 경우: 이미지에 4면 모두
        if has_caption:
            _set_cell_borders_sides(img_cell, {'left', 'right', 'top'},
                                    color=border_color, sz=6)
            _set_cell_borders_sides(cap_cell, {'left', 'right', 'bottom'},
                                    color=border_color, sz=6)
            _set_cell_vertical_padding(cap_cell, top_twips=10, bottom_twips=10)
        else:
            _set_cell_borders_sides(img_cell, {'left', 'right', 'top', 'bottom'},
                                    color=border_color, sz=6)

        # 이미지 삽입
        img_path = _resolve_image_path(img.get('url') or '', img.get('filename') or '')
        if img_path and os.path.exists(img_path):
            try:
                processed = _crop_to_aspect(img_path, 4/3)
                if processed != img_path:
                    temp_files.append(processed)
                p = img_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(processed,
                                width=Cm(img_width_cm),
                                height=Cm(img_height_cm))
            except Exception as e:
                p = img_cell.paragraphs[0]
                r = p.add_run(f'[이미지 로드 실패: {e}]')
                _set_font(r, size=9, color='B91C1C')
        else:
            p = img_cell.paragraphs[0]
            r = p.add_run('[이미지 없음]')
            _set_font(r, size=9, color='9CA3AF')

        # 캡션 (있는 경우만)
        if cap_cell is not None:
            cap_p = cap_cell.paragraphs[0]
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.space_before = Pt(0)
            cap_p.paragraph_format.space_after = Pt(0)
            cap_p.paragraph_format.line_spacing = 1.0
            caption = (img.get('caption') or '').strip()
            if caption:
                cap_run = cap_p.add_run(caption)
                _set_font(cap_run, size=9, color='4B5563')
                cap_run.italic = True

    # 마지막 그리드 행의 잉여 셀(이미지가 columns 수보다 적을 때) — 테두리 제거
    last_idx = n - 1
    last_gr = last_idx // columns
    last_col = last_idx % columns
    if last_col < columns - 1:
        img_row = row_offsets[last_gr]
        cap_row = img_row + 1 if row_has_caption[last_gr] else None
        for ci in range(last_col + 1, columns):
            _set_cell_borders_sides(tbl.rows[img_row].cells[ci], set())
            if cap_row is not None:
                _set_cell_borders_sides(tbl.rows[cap_row].cells[ci], set())

    col_cm_list = [cell_cm] * columns
    _set_table_fixed_layout(tbl, total_cm, col_cm_list)

    _GLOBAL_TEMP_FILES.extend(temp_files)

    _add_paragraph(doc, '', before=2, after=6)


# Word 생성 시 만들어진 임시 cropped 이미지 — build_docx 끝나면 일괄 삭제
_GLOBAL_TEMP_FILES = []


def _contain_in_static(candidate):
    """candidate 경로가 static 디렉터리 밖(경로순회 ../)이면 None 반환."""
    from app import app
    root = os.path.realpath(app.static_folder)
    real = os.path.realpath(candidate)
    if real == root or real.startswith(root + os.sep):
        return real
    return None


def _resolve_image_path(url, filename):
    """저장된 이미지 URL → 파일 시스템 경로 (static 디렉터리 밖 접근 차단)"""
    # url 예시: /static/uploads/dock/dock-1-1234-abc.jpg
    if url and url.startswith('/static/'):
        rel = url[len('/static/'):]
        # static 디렉터리 위치 — app 모듈에서 가져옴
        from app import app
        return _contain_in_static(os.path.join(app.static_folder, rel))
    # 직접 filename으로 fallback
    if filename:
        from app import app
        return _contain_in_static(
            os.path.join(app.static_folder, 'uploads', 'dock', filename))
    return None


# ─────────────────────────────────────────────────────────────
#  공개 함수 — 보고서 데이터 받아 docx 바이트 반환
# ─────────────────────────────────────────────────────────────
def build_docx(report: dict) -> bytes:
    """
    report: GET /api/dock-reports/<id> 응답과 동일한 구조
      · 메타: title, vessel_name, dock_no, shipyard, period_*, imo_no, gt, dwt,
              approval_*
      · sections: 평면 리스트, 각 항목에 blocks 포함
    """
    doc = Document()

    # 페이지 설정 (A4, 여백 2cm)
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width  = Mm(210)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # 기본 폰트
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    rFonts.set(qn('w:ascii'), 'Malgun Gothic')
    rFonts.set(qn('w:hAnsi'), 'Malgun Gothic')

    # 페이지 번호 (푸터 우측)
    _add_page_number_footer(doc)

    # 섹션 트리 빌드
    sections_flat = report.get('sections') or []
    tree = _build_tree(sections_flat)

    # ① 표지
    _build_cover(doc, report)
    # ② 목차
    if tree:
        _build_toc(doc, tree)
    # ③ 본문
    _render_sections(doc, tree)

    # 바이트로 직렬화
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    result = bio.read()

    # 이미지 cropped 임시 파일 정리
    global _GLOBAL_TEMP_FILES
    for fp in _GLOBAL_TEMP_FILES:
        try:
            os.remove(fp)
        except Exception:
            pass
    _GLOBAL_TEMP_FILES = []

    return result


def _build_tree(sections_flat):
    """평면 리스트 → parent_id 기반 트리"""
    by_id = {s['id']: dict(s, children=[]) for s in sections_flat}
    roots = []
    for s in by_id.values():
        if s.get('parent_id') and s['parent_id'] in by_id:
            by_id[s['parent_id']]['children'].append(s)
        else:
            roots.append(s)

    def sort_rec(lst):
        lst.sort(key=lambda x: (x.get('display_order', 0), x.get('id', 0)))
        for x in lst:
            sort_rec(x['children'])
    sort_rec(roots)
    return roots


def _add_page_number_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        # 페이지 번호 필드
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run._element.append(fld_begin)

        instr = OxmlElement('w:instrText')
        instr.text = 'PAGE'
        run._element.append(instr)

        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run._element.append(fld_end)

        _set_font(run, size=9, color='6B7280')
